import os
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        # Default chunking settings (can be overridden per KB)
        self.default_chunk_size = 500  # tokens per chunk
        self.default_chunk_overlap = 50  # overlap between chunks
        
    def determine_file_type(self, filename: str) -> str:
        """Determine file type from filename extension"""
        ext = Path(filename).suffix.lower()
        
        if ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
            return 'text'
        elif ext == '.pdf':
            return 'pdf'
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
            return 'image'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        elif ext == '.epub':
            return 'epub'
        else:
            return 'text'  # Default to text
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types"""
        try:
            if file_type == 'text':
                return self._extract_text_from_text_file(file_path)
            elif file_type == 'pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_type == 'image':
                return self._extract_text_from_image(file_path)
            elif file_type == 'docx':
                return self._extract_text_from_docx(file_path)
            elif file_type == 'epub':
                return self._extract_text_from_epub(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
                # If no text extracted (scanned PDF), try OCR
                if not text.strip():
                    logger.info(f"No text found in PDF {file_path}, attempting OCR")
                    # Note: This would require pdf2image for full OCR support
                    # For now, return empty string
                    return ""
                
                return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error performing OCR on {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_epub(self, file_path: str) -> str:
        """Extract text from EPUB files"""
        try:
            book = epub.read_epub(file_path)
            text = ""
            
            # Extract metadata
            title = book.get_metadata('DC', 'title')
            if title:
                text += f"Title: {title[0][0]}\n\n"
            
            creator = book.get_metadata('DC', 'creator')
            if creator:
                text += f"Author: {creator[0][0]}\n\n"
            
            # Extract text from all items
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Get the content
                    content = item.get_content().decode('utf-8')
                    
                    # Parse HTML content
                    soup = BeautifulSoup(content, 'html.parser')
                    
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # Get text and clean it up
                    chapter_text = soup.get_text()
                    
                    # Clean up whitespace
                    lines = (line.strip() for line in chapter_text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    chapter_text = ' '.join(chunk for chunk in chunks if chunk)
                    
                    if chapter_text.strip():
                        text += chapter_text + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error reading EPUB {file_path}: {str(e)}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = None, 
                   chunk_overlap: int = None, overlap_enabled: bool = True) -> List[str]:
        """Split text into chunks for vector storage
        
        Args:
            text: Text to chunk
            chunk_size: Size of chunks in tokens (uses default if None)
            chunk_overlap: Overlap between chunks in tokens (uses default if None)
            overlap_enabled: Whether to use overlapping chunks
        """
        if not text.strip():
            return []
        
        # Use provided values or defaults
        chunk_size = chunk_size or self.default_chunk_size
        chunk_overlap = chunk_overlap or self.default_chunk_overlap
        
        # Simple sentence-based chunking
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Rough token estimation (words * 1.3)
            sentence_tokens = len(sentence.split()) * 1.3
            current_tokens = len(current_chunk.split()) * 1.3
            
            if current_tokens + sentence_tokens > chunk_size and current_chunk:
                # Add current chunk
                chunks.append(current_chunk.strip())
                
                if overlap_enabled and chunk_overlap > 0:
                    # Create overlap by keeping last portion of current chunk
                    words = current_chunk.split()
                    overlap_words = int(chunk_overlap / 1.3)  # Convert tokens to words
                    overlap_words = min(overlap_words, len(words))
                    
                    if overlap_words > 0:
                        overlap_text = ' '.join(words[-overlap_words:])
                        current_chunk = overlap_text + ' ' + sentence + '. '
                    else:
                        current_chunk = sentence + '. '
                else:
                    # No overlap, start fresh
                    current_chunk = sentence + '. '
            else:
                current_chunk += sentence + '. '
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # Filter out very short chunks (at least 10 words)
        chunks = [chunk for chunk in chunks if len(chunk.split()) > 10]
        
        return chunks
    
    def process_file(self, file_path: str, filename: str, kb_id: str,
                    chunk_size: int = None, chunk_overlap: int = None,
                    overlap_enabled: bool = True) -> Dict[str, Any]:
        """Process a file and return document data for vector storage
        
        Args:
            file_path: Path to the file
            filename: Name of the file
            kb_id: Knowledge base ID
            chunk_size: Size of chunks in tokens (uses default if None)
            chunk_overlap: Overlap between chunks in tokens (uses default if None)
            overlap_enabled: Whether to use overlapping chunks
        """
        try:
            # Determine file type
            file_type = self.determine_file_type(filename)
            
            # Extract text
            text = self.extract_text_from_file(file_path, file_type)
            
            if not text.strip():
                logger.warning(f"No text extracted from {filename}")
                return {
                    "success": False,
                    "error": "No text content found in file"
                }
            
            # Chunk text with custom settings
            chunks = self.chunk_text(text, chunk_size, chunk_overlap, overlap_enabled)
            
            if not chunks:
                logger.warning(f"No chunks created from {filename}")
                return {
                    "success": False,
                    "error": "Could not create text chunks from file"
                }
            
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            return {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "file_type": file_type,
                "file_size": file_size,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"Processing failed: {str(e)}"
            }
    
    def save_file(self, file_content: bytes, filename: str, kb_id: str) -> str:
        """Save uploaded file to knowledge base directory"""
        try:
            # Create knowledge base directory - use absolute path from current working directory
            kb_dir = Path(f"knowledge-bases/{kb_id}")
            kb_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate unique filename to avoid conflicts
            file_ext = Path(filename).suffix
            file_stem = Path(filename).stem
            unique_filename = f"{file_stem}_{uuid.uuid4().hex[:8]}{file_ext}"
            
            file_path = kb_dir / unique_filename
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"Saved file {unique_filename} to {kb_dir}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error saving file {filename}: {str(e)}")
            raise
    
    def delete_file(self, file_path: str) -> bool:
        """Delete a file from the filesystem"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Deleted file {file_path}")
                return True
            else:
                logger.warning(f"File not found: {file_path}")
                return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {str(e)}")
            return False

# Global instance
file_processor = FileProcessor()
